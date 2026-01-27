"""
DOCX Text Extractor

Extracts text from DOCX files using python-docx.
"""

import logging
from typing import Optional
from io import BytesIO
from docx import Document

from ..interfaces.text_extraction_service import TextExtractionService

logger = logging.getLogger(__name__)


class DOCXExtractor(TextExtractionService):
    """DOCX text extraction using python-docx."""

    def supports_format(self, content_type: str, filename: Optional[str] = None) -> bool:
        """Check if file is a DOCX."""
        return (
            content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
            (filename and filename.lower().endswith('.docx'))
        )

    def extract_text(self, file_data: bytes, content_type: str, filename: Optional[str] = None) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_data: DOCX file data as bytes
            content_type: MIME type
            filename: Optional filename
            
        Returns:
            Extracted text from paragraphs and tables
            
        Raises:
            ValueError: If file is not a DOCX or is corrupted
        """
        if not self.supports_format(content_type, filename):
            raise ValueError(f"Unsupported format: {content_type}")
        
        try:
            # Open DOCX from bytes
            doc = Document(BytesIO(file_data))
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Join all parts with newlines
            full_text = '\n\n'.join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"DOCX {filename or 'unknown'} appears to be empty")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
